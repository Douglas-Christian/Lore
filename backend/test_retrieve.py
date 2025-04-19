import sys
import os
# Add the project root directory to the Python path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import requests
from pathlib import Path
import PyPDF2
import docx
import uuid
import chromadb
from chromadb.config import Settings
from backend.utils import debug_process_and_store_files
import time
import json

# Initialize ChromaDB client
chroma_client = chromadb.Client(Settings(persist_directory="chroma_data"))

def test_retrieve_endpoint():
    url = "http://127.0.0.1:8000/retrieve/"
    params = {"query": "leilon town council"}
    response = requests.get(url, params=params)
    print("Status Code:", response.status_code)
    print("Response JSON:", response.json())

def debug_indexing():
    """
    Debug the indexing process by extracting and printing text from sourcebooks.
    """
    path = Path("e:/Lore/backend/sourcebooks")
    files = path.glob("*.*")

    for file in files:
        extracted_text = ""
        if file.suffix.lower() == ".pdf":
            with open(file, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    extracted_text += page.extract_text()
        elif file.suffix.lower() == ".docx":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        elif file.suffix.lower() == ".txt":
            with open(file, "r", encoding="utf-8") as txt_file:
                extracted_text = txt_file.read()

        print(f"Extracted text from {file.name}:")
        print(extracted_text[:500])  # Print the first 500 characters for inspection

def debug_chromadb_indexing():
    """
    Debug the ChromaDB indexing process by printing the documents, metadata, and IDs being added.
    """
    path = Path("e:/Lore/backend/sourcebooks")
    files = path.glob("*.*")

    # Initialize ChromaDB collection
    collection = chroma_client.get_or_create_collection(name="dnd_sourcebooks")

    for file in files:
        extracted_text = ""
        if file.suffix.lower() == ".pdf":
            with open(file, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    extracted_text += page.extract_text()
        elif file.suffix.lower() == ".docx":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        elif file.suffix.lower() == ".txt":
            with open(file, "r", encoding="utf-8") as txt_file:
                extracted_text = txt_file.read()

        document_id = str(uuid.uuid4())
        print(f"Indexing document: {file.name}")
        print(f"Document ID: {document_id}")
        print(f"Metadata: {{'filename': {file.name}}}")
        print(f"Text snippet: {extracted_text[:500]}")

        if extracted_text.strip():
            collection.add(documents=[extracted_text], metadatas=[{"filename": file.name}], ids=[document_id])

def verify_full_text_extraction():
    """
    Verify if the full text of each page in the PDF is being extracted.
    """
    path = Path("e:/Lore/backend/sourcebooks/Storm Lords Wrath.pdf")
    if not path.exists():
        print(f"File '{path}' does not exist.")
        return

    extracted_text = ""
    with open(path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page_number, page in enumerate(reader.pages):
            page_text = page.extract_text()
            print(f"Page {page_number + 1} Text:")
            print(page_text[:500])  # Print the first 500 characters of each page for inspection
            extracted_text += page_text

    print("Total Extracted Text:")
    print(extracted_text[:1000])  # Print the first 1000 characters of the full text for inspection

def check_server_status():
    """Check if the FastAPI server is running"""
    try:
        response = requests.get("http://127.0.0.1:8000/")
        return response.status_code == 200
    except requests.exceptions.ConnectionError:
        return False

def test_narration_context_continuity():
    """
    Test that narration logs are properly stored and can be retrieved with context continuity.
    This function tests:
    1. Creating a campaign
    2. Adding multiple narration logs
    3. Verifying logs are retrievable via ChromaDB
    4. Testing context continuity across multiple narration entries
    """
    if not check_server_status():
        print("ERROR: FastAPI server is not running. Please start the server with 'uvicorn backend.main:app --reload'")
        return
    
    print("\n===== TESTING NARRATION CONTEXT CONTINUITY =====\n")
    
    # Step 1: Create a new campaign
    campaign_name = f"Narration Test Campaign {int(time.time())}"
    campaign_url = "http://127.0.0.1:8000/campaigns/"
    campaign_data = {"name": campaign_name, "description": "A campaign for testing narration context continuity."}
    
    print(f"Creating campaign: {campaign_name}")
    response = requests.post(campaign_url, params=campaign_data)
    
    if response.status_code != 200:
        print(f"Failed to create campaign: {response.text}")
        return
    
    campaign = response.json()
    campaign_id = campaign["id"]
    print(f"Created campaign with ID: {campaign_id}")
    
    # Step 2: Add first narration log
    first_narration = "The brave knight Sir Galahad arrived at the small village of Meadowbrook. The villagers welcomed him warmly, offering food and shelter. The village elder spoke of a terrible dragon terrorizing the countryside."
    narration_url = f"http://127.0.0.1:8000/campaigns/{campaign_id}/narration_logs/"
    narration_data = {"content": first_narration}
    
    print("\nAdding first narration log...")
    response = requests.post(narration_url, params=narration_data)
    
    if response.status_code != 200:
        print(f"Failed to add first narration log: {response.text}")
        return
    
    first_log = response.json()
    print(f"Added first narration log with ID: {first_log['id']}")
    
    # Step 3: Verify the first narration log is retrievable
    print("\nVerifying first narration log is retrievable...")
    search_terms = ["brave knight", "Galahad", "Meadowbrook", "dragon"]
    
    for term in search_terms:
        retrieve_url = "http://127.0.0.1:8000/retrieve/"
        params = {"query": term}
        response = requests.get(retrieve_url, params=params)
        
        if response.status_code != 200:
            print(f"Failed to retrieve content for '{term}': {response.text}")
            continue
        
        results = response.json()
        found = False
        
        # Check if our narration is in the results
        if 'results' in results and 'documents' in results['results']:
            documents = results['results']['documents'][0]
            for doc in documents:
                if first_narration in doc:
                    found = True
                    break
        
        print(f"Search term '{term}': {'Found' if found else 'Not found'} in retrieval results")
    
    # Step 4: Add a second narration log that builds on the first
    second_narration = "Sir Galahad decided to help the village of Meadowbrook. He spent the day gathering information about the dragon's lair from the villagers. Armed with this knowledge, he prepared for his journey to the mountain cave where the dragon was said to dwell."
    narration_data = {"content": second_narration}
    
    print("\nAdding second narration log...")
    response = requests.post(narration_url, params=narration_data)
    
    if response.status_code != 200:
        print(f"Failed to add second narration log: {response.text}")
        return
    
    second_log = response.json()
    print(f"Added second narration log with ID: {second_log['id']}")
    
    # Step 5: Verify context continuity
    print("\nVerifying context continuity across narrations...")
    continuity_terms = ["Galahad mountain cave", "dragon village Meadowbrook", "knight journey"]
    
    for term in continuity_terms:
        retrieve_url = "http://127.0.0.1:8000/retrieve/"
        params = {"query": term}
        response = requests.get(retrieve_url, params=params)
        
        if response.status_code != 200:
            print(f"Failed to retrieve content for '{term}': {response.text}")
            continue
        
        results = response.json()
        first_found = False
        second_found = False
        
        # Check if both narrations are in the results
        if 'results' in results and 'documents' in results['results']:
            documents = results['results']['documents'][0]
            for doc in documents:
                if first_narration in doc:
                    first_found = True
                if second_narration in doc:
                    second_found = True
        
        print(f"Continuity term '{term}':")
        print(f"  First narration: {'Found' if first_found else 'Not found'}")
        print(f"  Second narration: {'Found' if second_found else 'Not found'}")
    
    print("\n===== NARRATION CONTEXT CONTINUITY TEST COMPLETED =====\n")

if __name__ == "__main__":
    # Comment out or remove the previous test function call
    # test_narration_logging_and_retrieval()
    
    # Run the new, more comprehensive test
    test_narration_context_continuity()