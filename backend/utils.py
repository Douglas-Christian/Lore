import http.client as http_client
http_client.HTTPConnection.debuglevel = 1

# Enable logging for requests
import logging
logging.basicConfig(level=logging.DEBUG)
logging.getLogger("urllib3").setLevel(logging.DEBUG)

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.exc import NoResultFound
from backend.models import Campaign, NarrationLog, Session
import httpx
import json
from typing import Union
from pathlib import Path
import PyPDF2
import docx
import chromadb
from chromadb.config import Settings
import uuid

# Initialize ChromaDB client without telemetry
chroma_client = chromadb.Client(Settings(persist_directory="chroma_data"))

# Utility function to create a new campaign
async def create_campaign(db: AsyncSession, name: str, description: str = None):
    new_campaign = Campaign(name=name, description=description)
    db.add(new_campaign)
    await db.commit()
    await db.refresh(new_campaign)
    return new_campaign

# Utility function to retrieve all campaigns
async def get_campaigns(db: AsyncSession):
    result = await db.execute(select(Campaign))
    return result.scalars().all()

# Utility function to retrieve a campaign by ID
async def get_campaign_by_id(db: AsyncSession, campaign_id: int):
    try:
        result = await db.execute(select(Campaign).where(Campaign.id == campaign_id))
        return result.scalar_one()
    except NoResultFound:
        return None

# Utility function to create a new narration log
async def create_narration_log(db: AsyncSession, campaign_id: int, content: str):
    new_log = NarrationLog(campaign_id=campaign_id, content=content)
    db.add(new_log)
    await db.commit()
    await db.refresh(new_log)    
    return new_log

# Utility function to retrieve narration logs for a campaign
async def get_narration_logs(db: AsyncSession, campaign_id: int):
    result = await db.execute(select(NarrationLog).where(NarrationLog.campaign_id == campaign_id))
    return result.scalars().all()

# Utility function to create a new session
async def create_session(db: AsyncSession, campaign_id: int):
    new_session = Session(campaign_id=campaign_id)
    db.add(new_session)
    await db.commit()
    await db.refresh(new_session)
    return new_session

# Utility function to retrieve sessions for a campaign
async def get_sessions(db: AsyncSession, campaign_id: int):
    result = await db.execute(select(Session).where(Session.campaign_id == campaign_id))
    return result.scalars().all()

# Update the query_ollama function to handle streaming responses
def query_ollama(prompt: str, model: str = "llama3.2", api_key: str = None, timeout: int = 60) -> dict:
    """
    Query the Ollama Llama API with a given prompt using the /chat endpoint.

    Args:
        prompt (str): The input prompt for the LLM.
        model (str): The model to use (default is "llama3.2").
        api_key (str, optional): The API key for authentication. Defaults to None.
        timeout (int, optional): Timeout in seconds for the request. Defaults to 60.

    Returns:
        dict: The combined response from the LLM API.
    """
    url = "http://localhost:11434/api/chat"
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    payload = {
        "model": model,
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    print(f"Querying Ollama API at {url} with payload: {payload}")  # Debugging: Log the URL and payload

    try:
        with httpx.Client(timeout=timeout) as client:
            try:
                with client.stream("POST", url, json=payload, headers=headers) as response:
                    response.raise_for_status()
                    combined_response = ""
                    for chunk in response.iter_text():
                        print(f"Received chunk: {chunk}")  # Debugging: Log each chunk
                        # Extract the "content" field from each JSON object
                        try:
                            chunk_json = json.loads(chunk)
                            if "message" in chunk_json and "content" in chunk_json["message"]:
                                combined_response += chunk_json["message"]["content"]
                        except json.JSONDecodeError:
                            continue
                    
                    # If we've received no content, provide a fallback response
                    if not combined_response:
                        return {"response": "I couldn't generate a response based on the given context. Please try asking another question."}
                    
                    return {"response": combined_response}
            except httpx.TimeoutException:
                print("Request to Ollama API timed out")
                return {"error": "timed out", "fallback_response": "The request to the LLM timed out. This might be due to high server load or the complexity of the prompt. Consider using a simpler query or trying again later."}
    except httpx.RequestError as e:
        print(f"Error querying Ollama API: {e}")  # Debugging: Log the error
        return {"error": str(e), "fallback_response": "There was an error connecting to the LLM service. Please check that Ollama is running and try again."}

def process_and_store_files(path: Union[str, Path]):
    """
    Process a single file or all files in a directory and store their content in ChromaDB.

    Args:
        path (Union[str, Path]): Path to a file or directory.

    Returns:
        dict: A summary of processed files.
    """
    path = Path(path)
    if not path.exists():
        return {"error": f"Path '{path}' does not exist."}

    # Initialize ChromaDB collection
    collection = chroma_client.get_or_create_collection(name="dnd_sourcebooks")

    processed_files = []

    # Process a single file or all files in a directory
    files = [path] if path.is_file() else path.glob("*.*")

    for file in files:
        extracted_text = ""
        if file.suffix.lower() == ".pdf":
            with open(file, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    extracted_text += page.extract_text()
        elif file.suffix.lower() == ".docx":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        elif file.suffix.lower() == ".txt":
            with open(file, "r", encoding="utf-8") as txt_file:
                extracted_text = txt_file.read()
        else:
            continue  # Skip unsupported file types

        # Generate a unique ID for the document
        document_id = str(uuid.uuid4())

        # Store extracted text in ChromaDB
        if extracted_text.strip():
            collection.add(documents=[extracted_text], metadatas=[{"filename": file.name}], ids=[document_id])
            processed_files.append(file.name)

    return {"processed_files": processed_files}

def debug_process_and_store_files(path: Union[str, Path]):
    """
    Debug the process_and_store_files function by printing extracted text and verifying ChromaDB operations.
    """
    path = Path(path)
    if not path.exists():
        print(f"Path '{path}' does not exist.")
        return

    # Initialize ChromaDB collection
    collection = chroma_client.get_or_create_collection(name="dnd_sourcebooks")

    files = [path] if path.is_file() else path.glob("*.*")

    for file in files:
        extracted_text = ""
        if file.suffix.lower() == ".pdf":
            with open(file, "rb") as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                for page in reader.pages:
                    extracted_text += page.extract_text()
        elif file.suffix.lower() == ".docx":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
        elif file.suffix.lower() == ".txt":
            with open(file, "r", encoding="utf-8") as txt_file:
                extracted_text = txt_file.read()

        print(f"Extracted text from {file.name}:")
        print(extracted_text[:500])  # Print the first 500 characters for inspection

        # Generate a unique ID for the document
        document_id = str(uuid.uuid4())

        # Attempt to store extracted text in ChromaDB
        try:
            if extracted_text.strip():
                collection.add(documents=[extracted_text], metadatas=[{"filename": file.name}], ids=[document_id])
                print(f"Successfully added {file.name} to ChromaDB with ID {document_id}.")
            else:
                print(f"No text extracted from {file.name}, skipping.")
        except Exception as e:
            print(f"Error adding {file.name} to ChromaDB: {e}")

def retrieve_from_chromadb(query: str):
    """
    Retrieve content from the ChromaDB collection based on a query.

    Args:
        query (str): The search query.

    Returns:
        list: A list of matching documents and their metadata.
    """
    collection = chroma_client.get_or_create_collection(name="dnd_sourcebooks")
    results = collection.query(query_texts=[query], n_results=5)
    return results